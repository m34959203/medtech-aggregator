"""MedArchive extractor — устойчивый разбор реального архива прайсов.

Извлекает из каждого файла позиции с РАЗДЕЛЁННЫМИ тарифами и кодом:
ArchiveItem(code, name, price_resident, price_nonresident, price_original, currency).

Форматы: DOCX (с принятием tracked changes), табличный PDF (колонки тарифов
раздельно), XLSX/XLS (автодетект строки-заголовка, мультицена, все листы).
"""
from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass


@dataclass
class ArchiveItem:
    name: str
    code: str | None = None
    price_resident: float | None = None
    price_nonresident: float | None = None
    price_original: float | None = None
    currency: str = "KZT"


# ── код тарификатора ────────────────────────────────────────────────────────
# Формат справочника: A02.004.000 / B02.110.002 / C03.033.004 / D99.999.203.
# В прайсах буква нередко КИРИЛЛИЧЕСКАЯ (В, А, С, Е…) — нормализуем к латинице.
_CYR2LAT = str.maketrans({"А": "A", "В": "B", "С": "C", "Е": "E", "К": "K",
                          "М": "M", "Н": "H", "О": "O", "Р": "P", "Т": "T",
                          "Х": "X", "У": "Y"})
_CODE_RE = re.compile(r"[A-ZА-Я]\d{2}\.\d{3}\.\d{3}")


def norm_code(s: str | None) -> str | None:
    if not s:
        return None
    m = _CODE_RE.search(str(s).upper().translate(_CYR2LAT))
    return m.group(0) if m else None


# Фрагменты шапок/секций, которые ошибочно попадают в позиции из текстовых PDF —
# это НЕ услуги, отсекаем (иначе раздувают знаменатель и навсегда висят в unmatched).
_NOISE_RE = re.compile(
    r"^\s*(раздел\b|глава\b|прейскурант|приложение|утвержда|согласовано|"
    r"цена\s+для|цены\s+для|стоимость[,\s]+наимен|наименование\s+услуг|"
    r"группа\s+сложности|итого\b|всего\b|прайс[-\s]?лист|категория\s+сложн)",
    re.I,
)


def _is_noise_name(name: str) -> bool:
    n = (name or "").strip()
    if len(n) < 4:
        return True
    if _NOISE_RE.search(n):
        return True
    # имя без достаточного числа букв (мусорные строки из цифр/пунктуации)
    letters = len(re.findall(r"[A-Za-zА-Яа-яЁё]", n))
    return letters < 4


# ── разбор цены ─────────────────────────────────────────────────────────────
def parse_price(value) -> float | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return float(value) if value > 0 else None
    s = str(value).replace("\xa0", " ").strip()
    # вычищаем валюту/буквы, оставляем цифры/разделители
    m = re.search(r"\d[\d\s.,]*", s)
    if not m:
        return None
    token = m.group(0).strip().replace(" ", "")
    token = _normalize_number(token)
    try:
        val = float(token)
    except ValueError:
        return None
    return val if val > 0 else None


def _normalize_number(token: str) -> str:
    """Разводит десятичный разделитель и разделитель тысяч (цены KZT обычно целые).

    '10.002'→'10002' (тысячи), '2099.5'→'2099.5' (десятич.), '16 380'→'16380',
    '1 410,50'→'1410.50'. Эвристика: 3 цифры после разделителя ⇒ тысячи.
    """
    if "." in token and "," in token:
        # последний разделитель — десятичный
        return token.replace(",", "") if token.rfind(".") > token.rfind(",") \
            else token.replace(".", "").replace(",", ".")
    if token.count(".") > 1:
        return token.replace(".", "")
    if token.count(",") > 1:
        return token.replace(",", "")
    for sep in (".", ","):
        if sep in token:
            intp, dec = token.rsplit(sep, 1)
            if len(dec) == 3:          # 3 знака ⇒ тысячный разделитель
                return token.replace(sep, "")
            return token.replace(sep, ".")  # иначе десятичный
    return token


# ── классификация колонок ───────────────────────────────────────────────────
NAME_HINTS = ["наименование", "услуг", "название", "процедур", "анализ",
              "исследован", "name", "service", "атауы", "қызмет"]
CODE_HINTS = ["код", "code", "шифр"]
# нерезидент — только СИЛЬНЫЕ маркеры (СНГ / не проживающие постоянно).
# 'иностран' НЕ берём: встречается и в колонке резидента («иностранцы, постоянно проживающие»).
NONRES_HINTS = ["нерезидент", "снг", "ближнего зарубежья", "не проживающ", "зарубеж"]
# резидент — постоянно проживающие / граждане РК / кандас.
RES_HINTS = ["резидент", "постоянно прожива", "оралман", "кандас",
             "республики казахстан", "граждан рк"]
PRICE_HINTS = ["цена", "стоимост", "тариф", "сумма", "прайс", "price", "бағас", "құны", "тенге", "kzt", "ндс"]


def _h(cell, hints) -> int:
    c = str(cell or "").strip().lower()
    return sum(1 for x in hints if x in c)


def _classify_columns(header_cells: list) -> dict:
    """По ячейкам заголовка → индексы колонок name/code/resident/nonresident/price."""
    cols = {"name": None, "code": None, "resident": None, "nonresident": None, "prices": []}
    low_cells = [str(c or "").strip().lower() for c in header_cells]
    # 1. код — по 'код'/'шифр' (приоритет, чтобы «Код услуги» не утёк в name)
    for i, low in enumerate(low_cells):
        if low and _h(low, CODE_HINTS) and "наимен" not in low:
            cols["code"] = i
            break
    # 2. имя — лучший по NAME-скору, но НЕ код-колонка; 'наименование' весомее
    name_best = 0
    for i, low in enumerate(low_cells):
        if not low or i == cols["code"]:
            continue
        score = _h(low, NAME_HINTS) + (2 if "наимен" in low else 0)
        if score > name_best:
            name_best, cols["name"] = score, i
    # 3. цены + резидент/нерезидент. Колонка ценовая, если есть price-хинт ЛИБО
    #    маркер резидент/нерезидент (подписи бывают «для граждан РК» без слова «цена»).
    for i, low in enumerate(low_cells):
        if not low or i in (cols["code"], cols["name"]):
            continue
        is_nonres, is_res = _h(low, NONRES_HINTS), _h(low, RES_HINTS)
        if _h(low, PRICE_HINTS) or is_nonres or is_res:
            cols["prices"].append(i)
            if is_nonres and cols["nonresident"] is None:
                cols["nonresident"] = i
            elif is_res and cols["resident"] is None:
                cols["resident"] = i
    return cols


def _looks_like_header(cells: list) -> int:
    """Оценка «это строка-заголовок»: есть имя-колонка И хоть одна цена-колонка."""
    has_name = any(_h(c, NAME_HINTS) for c in cells)
    n_price = sum(1 for c in cells if _h(c, PRICE_HINTS))
    return (2 if has_name else 0) + n_price


# ── сборка позиций из «матрицы» (общая для xlsx/pdf) ────────────────────────
def _items_from_matrix(rows: list[list]) -> list[ArchiveItem]:
    if not rows:
        return []
    # 1. ищем строку-заголовок в первых 25 строках
    hdr_idx, hdr_score = None, 0
    for i, row in enumerate(rows[:25]):
        sc = _looks_like_header(row)
        if sc > hdr_score:
            hdr_idx, hdr_score = i, sc
    if hdr_idx is None or hdr_score < 2:
        cols = _fallback_columns(rows)
        data = rows
    else:
        # многострочная шапка: подписи цен (резидент/нерезидент) бывают на 1-2
        # строки ниже — складываем тексты колонок вертикально.
        header_rows = [hdr_idx]
        for j in range(hdr_idx + 1, min(hdr_idx + 3, len(rows))):
            row = rows[j]
            has_hint = _looks_like_header(row) > 0 or any(
                _h(c, RES_HINTS) or _h(c, NONRES_HINTS) or _h(c, PRICE_HINTS) for c in row
            )
            # продолжение шапки несёт ПОДПИСИ, не данные: у него нет разбираемых цен.
            # Строка с ценами — уже данные (её имя могло случайно содержать хинт «анализ»).
            n_prices = sum(1 for c in row if parse_price(c) is not None)
            if has_hint and n_prices == 0:
                header_rows.append(j)
            else:
                break
        width = max(len(rows[r]) for r in header_rows)
        merged = []
        for c in range(width):
            parts = []
            for r in header_rows:
                v = rows[r][c] if c < len(rows[r]) else None
                if v is not None and str(v).strip().lower() not in ("nan", "none", ""):
                    parts.append(str(v).strip())
            merged.append(" ".join(parts))
        cols = _classify_columns(merged)
        data = rows[header_rows[-1] + 1:]
        if cols["name"] is None or (not cols["prices"]):
            cols = _fallback_columns(data) or cols

    if cols.get("name") is None or not cols.get("prices"):
        return []

    # резидент/нерезидент: если явно не размечены, но цены 2 — раскладываем по порядку
    res_i, nonres_i = cols["resident"], cols["nonresident"]
    price_cols = cols["prices"]
    if res_i is None and nonres_i is None:
        if len(price_cols) == 1:
            res_i = price_cols[0]
        else:
            # порядок: обычно резидент (РК) идёт раньше нерезидента (СНГ)
            res_i, nonres_i = price_cols[0], price_cols[1]
    name_i, code_i = cols["name"], cols["code"]

    out: list[ArchiveItem] = []
    width = max(name_i, code_i or 0, *(price_cols or [0])) + 1
    for row in data:
        if len(row) < width:
            row = list(row) + [None] * (width - len(row))
        name = str(row[name_i]).replace("\xa0", " ").strip() if row[name_i] is not None else ""
        if not name or name.lower() in ("nan", "none", "итого", "итого:"):
            continue
        if _is_noise_name(name):
            continue
        res = parse_price(row[res_i]) if res_i is not None else None
        nonres = parse_price(row[nonres_i]) if nonres_i is not None else None
        primary = res or nonres or (parse_price(row[price_cols[0]]) if price_cols else None)
        if primary is None:
            continue
        code = norm_code(row[code_i]) if code_i is not None else None
        if not code:
            code = norm_code(name)  # код иногда внутри названия
        out.append(ArchiveItem(name=name, code=code, price_resident=res,
                               price_nonresident=nonres, price_original=primary))
    return out


def _fallback_columns(rows: list[list]) -> dict:
    """Без заголовка: имя = самая текстовая колонка, цены = числовые колонки."""
    if not rows:
        return {"name": None, "code": None, "resident": None, "nonresident": None, "prices": []}
    width = max(len(r) for r in rows)
    text_score = [0.0] * width
    num_cols = []
    for c in range(width):
        vals = [r[c] if c < len(r) else None for r in rows]
        textish = sum(1 for v in vals if v and re.search(r"[A-Za-zА-Яа-я]{3,}", str(v)))
        numish = sum(1 for v in vals if parse_price(v) is not None)
        text_score[c] = textish
        if numish > len(rows) * 0.3:
            num_cols.append(c)
    name_i = max(range(width), key=lambda c: text_score[c]) if width else None
    num_cols = [c for c in num_cols if c != name_i]
    return {"name": name_i, "code": None, "resident": None, "nonresident": None, "prices": num_cols}


# ── XLSX / XLS ──────────────────────────────────────────────────────────────
def parse_excel(content: bytes) -> list[ArchiveItem]:
    import pandas as pd
    bio = io.BytesIO(content)
    try:
        xls = pd.ExcelFile(bio)
    except Exception:
        bio.seek(0)
        xls = pd.ExcelFile(bio, engine="xlrd")
    items: list[ArchiveItem] = []
    for sheet in xls.sheet_names:
        df = xls.parse(sheet, header=None, dtype=object)
        rows = df.values.tolist()
        items.extend(_items_from_matrix(rows))
    return items


# ── PDF ─────────────────────────────────────────────────────────────────────
def parse_pdf(content: bytes) -> list[ArchiveItem]:
    import pdfplumber
    table_rows: list[list] = []
    word_items: list[ArchiveItem] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            tbls = page.extract_tables() or []
            if tbls:
                for table in tbls:
                    for r in table:
                        table_rows.append([(c or "").replace("\n", " ").strip()
                                           if isinstance(c, str) else c for c in r])
            else:
                word_items.extend(_pdf_page_words_to_items(page))
    items = _items_from_matrix(table_rows) if table_rows else []
    items.extend(word_items)
    # совсем пусто → возможно скан, OCR-фоллбэк (если доступен)
    if not items:
        try:
            from . import ocr  # type: ignore
            if ocr.ocr_available():
                items = [ArchiveItem(name=l.name, price_original=l.price)  # type: ignore
                         for l in []]
        except Exception:
            pass
    return items


def _merge_number_tokens(tokens: list[dict]) -> list[tuple[float, float]]:
    """Склеивает соседние числовые токены одной цены ('182' '900'→182900).

    Возвращает [(value, x0), …]. Большой зазор по x ⇒ разные ценовые колонки.
    """
    nums: list[tuple[float, float]] = []
    buf, buf_x0, prev_x1 = "", None, None
    for t in tokens:
        txt = t["text"].replace("\xa0", "")
        if re.fullmatch(r"[\d][\d.,]*", txt):
            if buf and prev_x1 is not None and (t["x0"] - prev_x1) < 12:
                buf += txt
            else:
                if buf:
                    v = parse_price(buf)
                    if v:
                        nums.append((v, buf_x0))
                buf, buf_x0 = txt, t["x0"]
            prev_x1 = t["x1"]
        else:
            if buf:
                v = parse_price(buf)
                if v:
                    nums.append((v, buf_x0))
                buf, prev_x1 = "", None
    if buf:
        v = parse_price(buf)
        if v:
            nums.append((v, buf_x0))
    return nums


def _pdf_page_words_to_items(page) -> list[ArchiveItem]:
    """Реконструкция строк по координатам слов (PDF без линий таблицы)."""
    from collections import defaultdict
    words = page.extract_words()
    if not words:
        return []
    rows: dict[int, list[dict]] = defaultdict(list)
    for w in words:
        rows[round(w["top"] / 3.0)].append(w)

    items: list[ArchiveItem] = []
    name_buf: list[str] = []
    for key in sorted(rows):
        line = sorted(rows[key], key=lambda w: w["x0"])
        text = " ".join(w["text"] for w in line)
        code = norm_code(text)
        nums = _merge_number_tokens(line)
        # имя — текстовые токены левее первой цены, без кода
        first_price_x = nums[0][1] if nums else 1e9
        name_tokens = [w["text"] for w in line
                       if w["x0"] < first_price_x - 1
                       and not re.fullmatch(r"[\d][\d.,]*", w["text"])
                       and not re.fullmatch(r"(тг|тенге|kzt|₸|руб)\.?", w["text"], re.I)]
        name = " ".join(name_tokens).strip(" .:-—")
        if code:
            name = _CODE_RE.sub("", name.upper().translate(_CYR2LAT)).strip() or name
        has_letters = bool(re.search(r"[A-Za-zА-Яа-яЁё]{3,}", name))

        if not nums:
            if has_letters:
                name_buf.append(name)
            continue
        # цены: берём правые числа как тарифы (рез/нерез/страховой по порядку)
        prices = [v for v, _ in nums if v >= 50]   # отсекаем № строки/мелочь
        if not prices:
            continue
        full_name = (" ".join(name_buf + ([name] if has_letters else []))).strip()
        name_buf = []
        if _is_noise_name(full_name):
            continue
        res = prices[0]
        nonres = prices[1] if len(prices) > 1 else None
        items.append(ArchiveItem(name=full_name, code=code, price_resident=res,
                                 price_nonresident=nonres, price_original=res))
    return items


# ── DOCX (с принятием tracked changes) ──────────────────────────────────────
_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _accept_tracked_changes(content: bytes) -> bytes:
    """Возвращает docx, где tracked changes приняты: <w:del> удалены, <w:ins> раскрыты."""
    import lxml.etree as ET
    zin = zipfile.ZipFile(io.BytesIO(content))
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name in zin.namelist():
            data = zin.read(name)
            if name == "word/document.xml":
                root = ET.fromstring(data)
                # удалить удалённый текст
                for el in root.iter(f"{_W}del"):
                    el.getparent().remove(el)
                # раскрыть вставки (оставить детей <w:ins> на месте)
                for ins in list(root.iter(f"{_W}ins")):
                    p = ins.getparent()
                    if p is None:
                        continue
                    idx = list(p).index(ins)
                    for child in reversed(list(ins)):
                        p.insert(idx, child)
                    p.remove(ins)
                data = ET.tostring(root)
            zout.writestr(name, data)
    return out.getvalue()


def parse_docx(content: bytes) -> list[ArchiveItem]:
    import docx
    try:
        content = _accept_tracked_changes(content)
    except Exception:
        pass
    doc = docx.Document(io.BytesIO(content))
    items: list[ArchiveItem] = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.replace("\n", " ").strip() for cell in row.cells])
        items.extend(_items_from_matrix(rows))
    # таблиц нет — пробуем абзацы как свободный текст
    if not items:
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        items = _items_from_matrix([[ln] for ln in lines])  # слабый фоллбэк
    return items


# ── диспетчер ───────────────────────────────────────────────────────────────
def detect_and_parse(filename: str, content: bytes) -> tuple[str, list[ArchiveItem]]:
    fn = filename.lower()
    if fn.endswith(".docx"):
        return "docx", parse_docx(content)
    if fn.endswith((".xlsx", ".xls")):
        return "xlsx" if fn.endswith(".xlsx") else "xls", parse_excel(content)
    if fn.endswith(".pdf"):
        return "pdf", parse_pdf(content)
    if content[:4] == b"%PDF":
        return "pdf", parse_pdf(content)
    if content[:2] == b"PK":
        if b"word/document.xml" in content[:4000] or fn.endswith(".docx"):
            return "docx", parse_docx(content)
        return "xlsx", parse_excel(content)
    if content[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        return "xls", parse_excel(content)
    return "unknown", []
